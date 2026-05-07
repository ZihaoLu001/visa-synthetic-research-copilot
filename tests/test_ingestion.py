from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document

from synthetic_researcher.ingestion import extract_survey_text, supported_upload_types

ROOT = Path(__file__).resolve().parents[1]


def test_supported_upload_types_cover_visa_testing_formats():
    assert {"txt", "md", "pdf", "docx", "csv", "xlsx"}.issubset(set(supported_upload_types()))


def test_extract_text_file():
    extracted = extract_survey_text(
        "survey.txt",
        b"1. How relevant is this value proposition?\n2. What annual fee would be acceptable?",
    )
    assert extracted.file_type == "txt"
    assert "annual fee" in extracted.text
    assert extracted.char_count == len(extracted.text)


def test_extract_docx_file():
    document = Document()
    document.add_paragraph("1. Which benefit feels most valuable?")
    document.add_paragraph("2. What barrier would stop you?")
    payload = BytesIO()
    document.save(payload)

    extracted = extract_survey_text("interview_guide.docx", payload.getvalue())

    assert extracted.file_type == "docx"
    assert "benefit" in extracted.text
    assert "barrier" in extracted.text


def test_extract_xlsx_file():
    payload = BytesIO()
    dataframe = pd.DataFrame(
        {
            "question": [
                "How relevant is this value proposition?",
                "What annual fee in CHF would feel acceptable?",
            ],
            "construct": ["adoption", "price"],
        }
    )
    with pd.ExcelWriter(payload, engine="openpyxl") as writer:
        dataframe.to_excel(writer, index=False, sheet_name="Survey")

    extracted = extract_survey_text("marketing_research_survey.xlsx", payload.getvalue())

    assert extracted.file_type == "xlsx"
    assert "[Sheet: Survey]" in extracted.text
    assert "annual fee" in extracted.text


def test_extract_public_pdf_demo_file():
    payload = (ROOT / "demo" / "public_survey_uploads" / "federal_reserve_mobile_payments_excerpt.pdf").read_bytes()

    extracted = extract_survey_text("federal_reserve_mobile_payments_excerpt.pdf", payload)

    assert extracted.file_type == "pdf"
    assert "mobile payment" in extracted.text.lower()
    assert "annual fee in CHF" in extracted.text
    assert extracted.extraction_notes == ["Extracted text from 1 PDF page(s)."]
