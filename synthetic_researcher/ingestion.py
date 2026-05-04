from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pandas as pd


SUPPORTED_UPLOAD_TYPES = ["txt", "md", "pdf", "docx", "csv", "xlsx"]


class SurveyExtractionError(ValueError):
    """Raised when an uploaded survey file cannot be converted into text."""


@dataclass(frozen=True)
class ExtractedSurveyText:
    text: str
    file_name: str
    file_type: str
    char_count: int
    extraction_notes: list[str]

    def metadata(self) -> dict[str, object]:
        return {
            "source": "uploaded_file",
            "file_name": self.file_name,
            "file_type": self.file_type,
            "char_count": self.char_count,
            "extraction_notes": self.extraction_notes,
        }


def supported_upload_types() -> list[str]:
    return list(SUPPORTED_UPLOAD_TYPES)


def extract_survey_text(file_name: str, payload: bytes) -> ExtractedSurveyText:
    extension = Path(file_name).suffix.lower().lstrip(".")
    if extension not in SUPPORTED_UPLOAD_TYPES:
        raise SurveyExtractionError(
            f"Unsupported file type '.{extension}'. Supported types: {', '.join(SUPPORTED_UPLOAD_TYPES)}."
        )

    notes: list[str] = []
    if extension in {"txt", "md"}:
        text = _decode_text(payload)
        notes.append("Decoded text directly from UTF-8 compatible content.")
    elif extension == "pdf":
        text, pages = _extract_pdf(payload)
        notes.append(f"Extracted text from {pages} PDF page(s).")
    elif extension == "docx":
        text, paragraphs, tables = _extract_docx(payload)
        notes.append(f"Extracted {paragraphs} paragraph(s) and {tables} table(s) from DOCX.")
    elif extension == "csv":
        text, rows = _extract_csv(payload)
        notes.append(f"Converted {rows} CSV row(s) into survey text.")
    elif extension == "xlsx":
        text, sheets, rows = _extract_xlsx(payload)
        notes.append(f"Converted {rows} row(s) from {sheets} Excel sheet(s) into survey text.")
    else:
        raise SurveyExtractionError(f"Unsupported file type '.{extension}'.")

    cleaned = _clean_text(text)
    if not cleaned:
        raise SurveyExtractionError("The uploaded file did not contain extractable survey text.")

    return ExtractedSurveyText(
        text=cleaned,
        file_name=file_name,
        file_type=extension,
        char_count=len(cleaned),
        extraction_notes=notes,
    )


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _extract_pdf(payload: bytes) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared in requirements
        raise SurveyExtractionError("PDF extraction requires the 'pypdf' package.") from exc

    reader = PdfReader(BytesIO(payload))
    page_text = []
    for page in reader.pages:
        page_text.append(page.extract_text() or "")
    return "\n\n".join(page_text), len(reader.pages)


def _extract_docx(payload: bytes) -> tuple[str, int, int]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared in requirements
        raise SurveyExtractionError("DOCX extraction requires the 'python-docx' package.") from exc

    document = Document(BytesIO(payload))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_count = 0
    for table in document.tables:
        table_count += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts), len(document.paragraphs), table_count


def _extract_csv(payload: bytes) -> tuple[str, int]:
    try:
        dataframe = pd.read_csv(BytesIO(payload))
    except Exception:
        return _decode_text(payload), 0
    return _dataframe_to_text(dataframe), len(dataframe)


def _extract_xlsx(payload: bytes) -> tuple[str, int, int]:
    try:
        sheets = pd.read_excel(BytesIO(payload), sheet_name=None, engine="openpyxl")
    except ImportError as exc:  # pragma: no cover - dependency is declared in requirements
        raise SurveyExtractionError("Excel extraction requires the 'openpyxl' package.") from exc
    except Exception as exc:
        raise SurveyExtractionError(f"Could not read Excel workbook: {exc}") from exc

    lines: list[str] = []
    total_rows = 0
    for sheet_name, dataframe in sheets.items():
        if dataframe.empty:
            continue
        total_rows += len(dataframe)
        lines.append(f"[Sheet: {sheet_name}]")
        lines.append(_dataframe_to_text(dataframe))
    return "\n".join(lines), len(sheets), total_rows


def _dataframe_to_text(dataframe: pd.DataFrame) -> str:
    lines: list[str] = []
    for _, row in dataframe.iterrows():
        cells = []
        for value in row.tolist():
            if pd.isna(value):
                continue
            text = str(value).strip()
            if text:
                cells.append(text)
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line)
